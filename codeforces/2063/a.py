from docx import Document

# Criando o documento
doc = Document()
doc.add_heading('Soluções dos Exercícios', level=1)

# Exercício 1
doc.add_heading('1. Número total de amostras possíveis (amostragem sem reposição)', level=2)
doc.add_paragraph(
    "A fórmula para calcular o número de combinações possíveis ao selecionar k elementos de um conjunto com n elementos é:"
)
doc.add_paragraph("C(n, k) = n! / (k! * (n-k)!)", style="Quote")
doc.add_paragraph(
    "Neste caso, n = 100 (funcionários) e k = 25 (amostra de funcionários). Então:"
)
doc.add_paragraph(
    "C(100, 25) = 100! / (25! * 75!)", style="Quote"
)
doc.add_paragraph(
    "O resultado é um número extremamente grande, mas pode ser calculado de forma eficiente usando software ou calculadoras científicas. "
    "Para efeito de aproximação, o valor é cerca de:"
)
doc.add_paragraph("C(100, 25) ≈ 2.42 × 10^23", style="Quote")
doc.add_paragraph(
    "Portanto, o número total de amostras possíveis é aproximadamente 2.42 × 10^23."
)

# Exercício 2
doc.add_heading('2. Amostragem estratificada proporcional', level=2)
doc.add_paragraph(
    "Os tamanhos de cada categoria são:\n"
    "- Alunos: 1400\n"
    "- Professores: 150\n"
    "- Funcionários: 88\n"
)
doc.add_paragraph(
    "O total da população é: 1400 + 150 + 88 = 1638."
)
doc.add_paragraph(
    "Se a amostra total será de 70 pessoas, calcula-se a proporção de cada categoria na população e aplica-se a mesma proporção na amostra:"
)

doc.add_paragraph("1. Alunos:", style="List Number")
doc.add_paragraph(
    "Proporção de alunos = 1400 / 1638 ≈ 0.8549\n"
    "Quantidade de alunos na amostra = 0.8549 × 70 ≈ 59.84 ≈ 60"
)

doc.add_paragraph("2. Professores:", style="List Number")
doc.add_paragraph(
    "Proporção de professores = 150 / 1638 ≈ 0.0916\n"
    "Quantidade de professores na amostra = 0.0916 × 70 ≈ 6.41 ≈ 6"
)

doc.add_paragraph("3. Funcionários:", style="List Number")
doc.add_paragraph(
    "Proporção de funcionários = 88 / 1638 ≈ 0.0537\n"
    "Quantidade de funcionários na amostra = 0.0537 × 70 ≈ 3.76 ≈ 4"
)

doc.add_paragraph(
    "Assim, a amostra estratificada será composta por:\n"
    "- 60 alunos\n"
    "- 6 professores\n"
    "- 4 funcionários"
)

# Exercício 3
doc.add_heading('3. Amostra sistemática', level=2)
doc.add_paragraph(
    "Para obter uma amostra sistemática de 40 alunos a partir de uma população de 800, o procedimento é o seguinte:"
)
doc.add_paragraph("1. Calcular o intervalo de seleção (k):", style="List Number")
doc.add_paragraph(
    "k = População total / Tamanho da amostra = 800 / 40 = 20"
)
doc.add_paragraph("2. Escolher um ponto de partida inicial aleatório:", style="List Number")
doc.add_paragraph(
    "Selecione um número aleatório entre 1 e k = 20. Suponha que o número inicial escolhido seja 7."
)
doc.add_paragraph("3. Selecionar os alunos a cada k unidades:", style="List Number")
doc.add_paragraph(
    "Os alunos selecionados serão: 7, 27, 47, 67, ..., 787 (somando 20 sucessivamente até atingir 40 alunos)."
)

doc.add_paragraph(
    "Resumo do procedimento:\n"
    "- Calcular o intervalo (k = 20).\n"
    "- Escolher um ponto de partida aleatório (7, por exemplo).\n"
    "- Selecionar os alunos em saltos regulares de 20."
)

# Salvando o documento
file_path = "/mnt/data/Solucoes_Exercicios.docx"
doc.save(file_path)
file_path
